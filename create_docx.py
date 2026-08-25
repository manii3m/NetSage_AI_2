from docx import Document
from docx.shared import Pt, Inches

doc = Document()
# Title
title = doc.add_heading('NetSage AI - Intelligent Network Diagnostic Platform', 0)
title.alignment = 1

doc.add_paragraph('Team Members:').bold = True
doc.add_paragraph('• Mayank Vishwakarma (Father of CSE)')
doc.add_paragraph('• Deepak Mishra (Mother of CSE)')

doc.add_paragraph('College: Lakshmi Narain College Of Technology And Excellence Bhopal (M.P.)')
doc.add_paragraph('Technology Track: AI Track')

doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph('NetSage AI is an intelligent and automated network diagnostic platform designed for Cisco IOS and Cisco Packet Tracer environments. The main purpose of the project is to help network engineers and students identify and understand network faults by analyzing symptoms reported by the user, network topology details, and Cisco CLI command outputs.')
doc.add_paragraph('While working with network configurations, troubleshooting can often become time-consuming. An engineer may need to go through multiple "show" command outputs, check interface states, verify VLAN configurations, examine routing tables, and identify problems across different OSI layers. NetSage AI is designed to simplify this process by organizing the available network evidence and providing a structured diagnosis.')
doc.add_paragraph('The project follows a hybrid diagnostic approach. Instead of depending entirely on Artificial Intelligence, the system first uses deterministic, rule-based checks to identify known and common network problems. If a known pattern is detected, the system provides a direct diagnosis based on predefined logic. However, if the available evidence does not match any known rule, the case is passed to a Large Language Model (LLM) for further evidence-based analysis.')
doc.add_paragraph('This approach gives the project a good balance between reliability and flexibility. Known issues can be handled in a controlled and predictable way, while more complex or unusual problems can still benefit from AI-based reasoning. The final diagnosis is also reviewed by a human before any action is accepted, making the overall workflow safer and more practical for real-world network troubleshooting.')

doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph('Network troubleshooting is not always straightforward, especially in multi-layer Cisco network environments. A single connectivity issue can have several possible causes, such as an interface being administratively down, an incorrect IP configuration, a VLAN mismatch, missing routes, access control restrictions, or configuration errors at another layer of the network.')
doc.add_paragraph('Traditionally, the person troubleshooting the issue needs to manually run multiple Cisco CLI commands and carefully examine the output. This process requires networking knowledge and can take a significant amount of time, particularly for students or less experienced engineers. In many cases, the CLI output itself is lengthy, making it difficult to quickly identify the actual cause of the problem.')
doc.add_paragraph('Using a completely AI-based solution also creates certain risks. An AI model may make assumptions about information that was never provided or suggest commands that are not appropriate for the actual network configuration. On the other hand, a traditional rule-based system is reliable for known problems but cannot always handle unusual, ambiguous, or complex situations.')
doc.add_paragraph('NetSage AI was developed to address this gap. The system combines deterministic rule-based validation, LLM-based reasoning, and human review in one workflow. This ensures that known problems are handled predictably, while more complex cases can still receive intelligent analysis without giving the AI unrestricted control over the troubleshooting process.')

doc.add_heading('3. Proposed Solution', level=1)
doc.add_paragraph('NetSage AI provides an end-to-end workflow for diagnosing network faults using the evidence available from the network environment.')
doc.add_paragraph('The system receives three main inputs:').bold = True
doc.add_paragraph('Symptom: The problem reported by the user or network operator.', style='List Bullet')
doc.add_paragraph('Topology Note: Relevant information about the network structure, including devices, interfaces, connections, and other important topology details.', style='List Bullet')
doc.add_paragraph('Show Outputs: Cisco IOS CLI outputs that act as the main evidence for the diagnostic process.', style='List Bullet')
doc.add_paragraph('The first stage is handled by the deterministic rule engine in "checker.py". This component checks the available evidence against predefined patterns and rules for known network problems.')
doc.add_paragraph('If no rule matches the available evidence, the case is passed to "engine.py", which prepares the information for LLM-based analysis. The LLM is instructed to reason only from the evidence provided by the system and avoid inventing missing network details.')
doc.add_paragraph('The final output is returned in a structured format containing: Root cause of the problem, Affected OSI layer, Confidence score, Supporting evidence, Recommended verification command, and Suggested fix steps.')

doc.add_heading('4. System Architecture and Workflow', level=1)
doc.add_paragraph('The architecture is divided into four main areas:')
doc.add_paragraph('4.1 Data Tier: Contains the structured network test cases, case information, and system configuration required by the application.')
doc.add_paragraph('4.2 Diagnostic Core: The main processing part of the system. It includes the Deterministic rule checker, Main diagnostic engine, and LLM prompt and analysis logic.')
doc.add_paragraph('4.3 Human-in-the-Loop Gate: The diagnosis is displayed to the operator, who can review the result and decide whether to Approve, Edit, or Reject the diagnosis.')
doc.add_paragraph('4.4 Audit and Logging Layer: Records important information about the diagnostic process and the final human decision.')

doc.add_heading('5. Overall System Workflow', level=1)
doc.add_paragraph('1. The operator selects or provides a network diagnostic case.', style='List Number')
doc.add_paragraph('2. The system loads the symptom, topology information, and Cisco CLI outputs.', style='List Number')
doc.add_paragraph('3. The deterministic checker analyzes the evidence first.', style='List Number')
doc.add_paragraph('4. If a known issue is identified, a structured diagnosis is generated directly.', style='List Number')
doc.add_paragraph('5. If no rule matches, the diagnostic engine sends the available evidence to the LLM.', style='List Number')
doc.add_paragraph('6. The LLM returns an evidence-based diagnosis in the required format.', style='List Number')
doc.add_paragraph('7. The result is displayed on the dashboard for human review.', style='List Number')
doc.add_paragraph('8. The operator can approve, edit, or reject the proposed action.', style='List Number')
doc.add_paragraph('9. The final decision and diagnostic information are stored for auditing and future analysis.', style='List Number')

doc.add_heading('6. Technology Stack', level=1)
doc.add_paragraph('• Python: Core application logic, diagnostic processing, and integration.')
doc.add_paragraph('• Streamlit: Interactive dashboard where users can view network cases.')
doc.add_paragraph('• Pandas: Loading and processing structured network case data.')
doc.add_paragraph('• JSON: Structured format for communication between modules.')
doc.add_paragraph('• Cisco IOS CLI and Packet Tracer: Primary diagnostic environment.')
doc.add_paragraph('• Markdown and Mermaid.js: Project documentation and architecture.')

doc.add_heading('7. Team Roles and Responsibilities', level=1)
doc.add_paragraph('• Role 1: Project Lead - Overall direction, planning, coordination, and architecture design.')
doc.add_paragraph('• Role 2: Backend Developer - Building the core diagnostic engine, deterministic rules, and JSON handling.')
doc.add_paragraph('• Role 3: Integration Developer - Connecting major components, LLM integration, and Streamlit communication.')
doc.add_paragraph('• Role 4: Testing and Validation Developer - Ensuring reliable, valid, and safe diagnostic results.')

doc.add_heading('8. Key Feature: Human-in-the-Loop Validation', level=1)
doc.add_paragraph('One of the most important decisions in the design of NetSage AI is that an AI-generated recommendation is not automatically treated as the final action. The operator can approve, edit, or reject the diagnosis. This keeps the final decision under human control and adds an important safety layer to the system.')

doc.add_heading('9. Key Strengths of the Project', level=1)
doc.add_paragraph('• Deterministic Rules provide predictable and reliable handling of known network problems.')
doc.add_paragraph('• LLM-Based Reasoning provides flexibility for cases that cannot be handled by predefined rules alone.')
doc.add_paragraph('• Human Validation ensures that the final action remains under the control of a person.')
doc.add_paragraph('• Modular Architecture allows new diagnostic rules or capabilities to be added easily.')

doc.add_heading('10. Conclusion', level=1)
doc.add_paragraph('NetSage AI is a practical software project that brings together Python development, Cisco networking concepts, deterministic rule-based analysis, LLM integration, and human validation to address a real network troubleshooting problem. The project is intentionally designed not to depend completely on AI. Instead, it follows a controlled hybrid approach where known network faults are handled using deterministic rules first, and LLMs are used for additional reasoning when necessary.')

doc.save('Overview.docx')
print('Word document saved successfully.')
